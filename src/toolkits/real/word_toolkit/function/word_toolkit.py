import os
from docx import Document
from typing import Optional
from docx.shared import Inches, Pt, RGBColor


def word_create_document(file_path: str) -> dict:
    """
    Create a new Word document.

    Args:
        file_path (str): The path where the document will be saved.

    Returns:
        Document: A new Word document object.
    """
    dir_path = os.path.dirname(file_path)
    if dir_path and not os.path.exists(dir_path):
        os.makedirs(dir_path, exist_ok=True)
    doc = Document()
    doc.save(file_path)
    return {"success": True, "message": f"Document  {file_path} created successfully."}


def word_add_heading(file_path: str, text: str, level: int = 1) -> dict:
    """
    Add a heading to the Word document.

    Args:
        file_path (str): The path of the Word document.
        text (str): The text of the heading.
        level (int): The level of the heading (1-9). Default is 1.

    Returns:
        dict: A dictionary with success status and message.
    """
    if not os.path.exists(file_path):
        return {"success": False, "message": f"File {file_path} does not exist."}
    if not isinstance(level, int) or level < 1 or level > 9:
        return {"success": False, "message": "Level must be an integer between 1 and 9."}
    try:
        doc = Document(file_path)
        doc.add_heading(text, level=level)
        doc.save(file_path)
        return {"success": True, "message": f"Heading '{text}' added to {file_path}."}
    except Exception as e:
        return {"success": False, "message": f"Error adding heading: {e}"}
    

def word_add_paragraph(file_path: str, text: str) -> dict:
    """
    Add a paragraph to the Word document.

    Args:
        file_path (str): The path of the Word document.
        text (str): The text of the paragraph.

    Returns:
        dict: A dictionary with success status and message.
    """
    if not os.path.exists(file_path):
        return {"success": False, "message": f"File {file_path} does not exist."}
    try:
        doc = Document(file_path)
        doc.add_paragraph(text)
        doc.save(file_path)
        return {"success": True, "message": f"Paragraph '{text}' added to {file_path}."} 
    except Exception as e:
        return {"success": False, "message": f"Error adding paragraph: {e}"}
    

def word_add_table(file_path: str, rows: int, cols: int, data: list[list[str]]) -> dict:
    """
    Add a table to the Word document.

    Args:
        file_path (str): The path of the Word document.
        rows (int): The number of rows in the table.
        cols (int): The number of columns in the table.
        data (list[list[str]]): The data to fill the table. Each inner list represents a row.

    Returns:
        dict: A dictionary with success status and message.
    """
    if not os.path.exists(file_path):
        return {"success": False, "message": f"File {file_path} does not exist."}
    try:
        doc = Document(file_path)
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for i, row_data in enumerate(data):
            if i >= rows:
                break
            for j, cell_data in enumerate(row_data):
                if j >= cols:
                    break
                table.cell(i, j).text = str(cell_data)
        doc.save(file_path)
        return {"success": True, "message": f"Table added to {file_path}."}    
    except Exception as e:
        return {"success": False, "message": f"Error adding table: {e}"}
    

def word_delete_paragraph(file_path: str, paragraph_index: int) -> dict:
    """
    Delete a paragraph from the Word document.

    Args:
        file_path (str): The path of the Word document.
        paragraph_index (int): The index of the paragraph to delete.

    Returns:
        dict: A dictionary with success status and message.
    """
    if not os.path.exists(file_path):
        return {"success": False, "message": f"File {file_path} does not exist."}
    try:
        doc = Document(file_path)
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return {"success": False, "message": "Paragraph index out of range."}
        p = doc.paragraphs[paragraph_index]
        p._element.getparent().remove(p._element)
        doc.save(file_path)
        return {"success": True, "message": f"Paragraph at index {paragraph_index} deleted from {file_path}."}
    except Exception as e:
        return {"success": False, "message": f"Error deleting paragraph: {e}"}
    

def word_add_picture(file_path: str, image_path: str, width: Optional[float] = None, height: Optional[float] = None) -> dict:
    """
    Add a picture to the Word document.

    Args:
        file_path (str): The path of the Word document.
        image_path (str): The path of the image to add.
        width (float): The width of the image in inches. Default is None.
        height (float): The height of the image in inches. Default is None.

    Returns:
        dict: A dictionary with success status and message.
    """
    if not os.path.exists(file_path):
        return {"success": False, "message": f"File {file_path} does not exist."}
    if not os.path.exists(image_path):
        return {"success": False, "message": f"Image {image_path} does not exist."}
    try:
        doc = Document(file_path)
        doc.add_picture(image_path, width=width, height=height)
        doc.save(file_path)
        return {"success": True, "message": f"Picture added to {file_path}."}
    except Exception as e:
        return {"success": False, "message": f"Error adding picture: {e}"}
    

def word_add_page_break(file_path: str) -> dict:
    """
    Add a page break to the Word document.

    Args:
        file_path (str): The path of the Word document.

    Returns:
        dict: A dictionary with success status and message.
    """
    if not os.path.exists(file_path):
        return {"success": False, "message": f"File {file_path} does not exist."}
    try:
        doc = Document(file_path)
        doc.add_page_break()
        doc.save(file_path)
        return {"success": True, "message": f"Page break added to {file_path}."}
    except Exception as e:
        return {"success": False, "message": f"Error adding page break: {e}"}
    

def word_search_and_replace(file_path: str, search_text: str, replace_text: str) -> dict:
    """
    Search and replace text in the Word document.

    Args:
        file_path (str): The path of the Word document.
        search_text (str): The text to search for.
        replace_text (str): The text to replace with.

    Returns:
        dict: A dictionary with success status and message.
    """
    if not os.path.exists(file_path):
        return {"success": False, "message": f"File {file_path} does not exist."}
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if search_text in paragraph.text:
                paragraph.text = paragraph.text.replace(search_text, replace_text)
        doc.save(file_path)
        return {"success": True, "message": f"Text '{search_text}' replaced with '{replace_text}' in {file_path}."}
    except Exception as e:
        return {"success": False, "message": f"Error replacing text: {e}"}
    

def word_format_text(
        file_path: str, 
        text: str, 
        font_name: Optional[str] = "Times New Roman", 
        font_size: Optional[int] = 12, 
        bold: Optional[bool] = False, 
        italic: Optional[bool] = False, 
        underline: Optional[bool] = False, 
        color: Optional[str] = "000000"
    ) -> dict:
    """
    Format text in the Word document.

    Args:
        file_path (str): The path of the Word document.
        text (str): The text to format.
        font_name (str): The font name. Default is "Times New Roman".
        font_size (int): The font size. Default is 12.
        bold (bool): Whether to make the text bold. Default is False.
        italic (bool): Whether to make the text italic. Default is False.
        underline (bool): Whether to underline the text. Default is False.
        color (str): The color of the text in hex format. Default is "000000".

    Returns:
        dict: A dictionary with success status and message.
    """
    if not os.path.exists(file_path):
        return {"success": False, "message": f"File {file_path} does not exist."}
    try:
        doc = Document(file_path)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        run.font.color.rgb = RGBColor(int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16))
        doc.save(file_path)
        return {"success": True, "message": f"Text '{text}' formatted in {file_path}."}
    except Exception as e:
        return {"success": False, "message": f"Error formatting text: {e}"}
    

def word_get_document_info(file_path: str) -> dict:
    """
    Get information about the Word document.

    Args:
        file_path (str): The path of the Word document.

    Returns:
        dict: A dictionary with success status and document information.
    """
    if not os.path.exists(file_path):
        return {"success": False, "message": f"File {file_path} does not exist."}
    try:
        doc = Document(file_path)
        info = {
            "title": doc.core_properties.title,
            "author": doc.core_properties.author,
            "subject": doc.core_properties.subject,
            "keywords": doc.core_properties.keywords,
            "last_modified_by": doc.core_properties.last_modified_by,
            "created": str(doc.core_properties.created),
            "modified": str(doc.core_properties.modified),
            "category": doc.core_properties.category,
            "content_status": doc.core_properties.content_status
        }
        return {"success": True, "info": info}
    except Exception as e:
        return {"success": False, "message": f"Error getting document info: {e}"}
    

def word_get_document_text(file_path: str) -> dict:
    """
    Get the text content of the Word document.

    Args:
        file_path (str): The path of the Word document.

    Returns:
        dict: A dictionary with success status and document text.
    """
    if not os.path.exists(file_path):
        return {"success": False, "message": f"File {file_path} does not exist."}
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return {"success": True, "text": text}
    except Exception as e:
        return {"success": False, "message": f"Error getting document text: {e}"}


def word_get_document_outline(file_path: str) -> dict:
    """
    Get the outline of the Word document.

    Args:
        file_path (str): The path of the Word document.

    Returns:
        dict: A dictionary with success status and document outline.
    """
    if not os.path.exists(file_path):
        return {"success": False, "message": f"File {file_path} does not exist."}
    try:
        doc = Document(file_path)
        outline = []
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                outline.append(paragraph.text)
        return {"success": True, "outline": outline}
    except Exception as e:
        return {"success": False, "message": f"Error getting document outline: {e}"}
    

def word_get_document_paragraph(file_path: str, index: int) -> dict:
    """
    Get a specific paragraph from the Word document.

    Args:
        file_path (str): The path of the Word document.
        index (int): The index of the paragraph.

    Returns:
        dict: A dictionary with success status and paragraph text.
    """
    if not os.path.exists(file_path):
        return {"success": False, "message": f"File {file_path} does not exist."}
    try:
        doc = Document(file_path)
        if index < 0 or index >= len(doc.paragraphs):
            return {"success": False, "message": "Paragraph index out of range."}
        paragraph = doc.paragraphs[index]
        return {"success": True, "paragraph": paragraph.text}
    except Exception as e:
        return {"success": False, "message": f"Error getting paragraph: {e}"}
    

def word_copy_document(source_file_path: str, destination_file_path: str) -> dict:
    """
    Copy a Word document.

    Args:
        source_file_path (str): The path of the source Word document.
        destination_file_path (str): The path of the destination Word document.

    Returns:
        dict: A dictionary with success status and message.
    """
    if not os.path.exists(source_file_path):
        return {"success": False, "message": f"Source file {source_file_path} does not exist."}
    try:
        doc = Document(source_file_path)
        doc.save(destination_file_path)
        return {"success": True, "message": f"Document copied to {destination_file_path}."}
    except Exception as e:
        return {"success": False, "message": f"Error copying document: {e}"}
    

def word_list_documents(directory: str) -> dict:
    """
    List all Word documents in a directory.

    Args:
        directory (str): The path of the directory.

    Returns:
        dict: A dictionary with success status and list of document paths.
    """
    if not os.path.exists(directory):
        return {"success": False, "message": f"Directory {directory} does not exist."}
    try:
        files = [f for f in os.listdir(directory) if f.endswith('.docx')]
        file_paths = [os.path.join(directory, f) for f in files]
        return {"success": True, "documents": file_paths}
    except Exception as e:
        return {"success": False, "message": f"Error listing documents: {e}"}
    
