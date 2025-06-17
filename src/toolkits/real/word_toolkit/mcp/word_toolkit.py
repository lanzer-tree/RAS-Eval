from mcp.server.fastmcp import FastMCP
from src.toolkits.real.word_toolkit.function.word_toolkit import word_create_document
from src.toolkits.real.word_toolkit.function.word_toolkit import word_add_heading
from src.toolkits.real.word_toolkit.function.word_toolkit import word_add_paragraph
from src.toolkits.real.word_toolkit.function.word_toolkit import word_add_table
from src.toolkits.real.word_toolkit.function.word_toolkit import word_delete_paragraph
from src.toolkits.real.word_toolkit.function.word_toolkit import word_add_picture
from src.toolkits.real.word_toolkit.function.word_toolkit import word_add_page_break
from src.toolkits.real.word_toolkit.function.word_toolkit import word_search_and_replace
from src.toolkits.real.word_toolkit.function.word_toolkit import word_format_text
from src.toolkits.real.word_toolkit.function.word_toolkit import word_get_document_info
from src.toolkits.real.word_toolkit.function.word_toolkit import word_get_document_text
from src.toolkits.real.word_toolkit.function.word_toolkit import word_get_document_outline
from src.toolkits.real.word_toolkit.function.word_toolkit import word_get_document_paragraph
from src.toolkits.real.word_toolkit.function.word_toolkit import word_copy_document
from src.toolkits.real.word_toolkit.function.word_toolkit import word_list_documents
import os
from docx import Document
from typing import Optional
from docx.shared import Inches, Pt, RGBColor


mcp = FastMCP('word_toolkit')


@mcp.tool()
def _word_create_document(file_path: str) -> dict:
    """
    Create a new Word document.

    Args:
        file_path (str): The path where the document will be saved.

    Returns:
        Document: A new Word document object.
    """
    return word_create_document(file_path=file_path)


@mcp.tool()
def _word_add_heading(file_path: str, text: str, level: int=1) -> dict:
    """
    Add a heading to the Word document.

    Args:
        file_path (str): The path of the Word document.
        text (str): The text of the heading.
        level (int): The level of the heading (1-9). Default is 1.

    Returns:
        dict: A dictionary with success status and message.
    """
    return word_add_heading(file_path=file_path, text=text, level=level)


@mcp.tool()
def _word_add_paragraph(file_path: str, text: str) -> dict:
    """
    Add a paragraph to the Word document.

    Args:
        file_path (str): The path of the Word document.
        text (str): The text of the paragraph.

    Returns:
        dict: A dictionary with success status and message.
    """
    return word_add_paragraph(file_path=file_path, text=text)


@mcp.tool()
def _word_add_table(file_path: str, rows: int, cols: int, data: list[list[str]]) -> dict:
    """
    Add a table to the Word document.

    Args:
        file_path (str): The path of the Word document.
        rows (int): The number of rows in the table.
        cols (int): The number of columns in the table.
        data (list[list[str]]): The data to fill the table. Each inner list represents a row.

    Returns:
        dict: A dictionary with success status and message.
    """
    return word_add_table(file_path=file_path, rows=rows, cols=cols, data=data)


@mcp.tool()
def _word_delete_paragraph(file_path: str, paragraph_index: int) -> dict:
    """
    Delete a paragraph from the Word document.

    Args:
        file_path (str): The path of the Word document.
        paragraph_index (int): The index of the paragraph to delete.

    Returns:
        dict: A dictionary with success status and message.
    """
    return word_delete_paragraph(file_path=file_path, paragraph_index=paragraph_index)


@mcp.tool()
def _word_add_picture(file_path: str, image_path: str, width: Optional[float]=None, height: Optional[float]=None) -> dict:
    """
    Add a picture to the Word document.

    Args:
        file_path (str): The path of the Word document.
        image_path (str): The path of the image to add.
        width (float): The width of the image in inches. Default is None.
        height (float): The height of the image in inches. Default is None.

    Returns:
        dict: A dictionary with success status and message.
    """
    return word_add_picture(file_path=file_path, image_path=image_path, width=width, height=height)


@mcp.tool()
def _word_add_page_break(file_path: str) -> dict:
    """
    Add a page break to the Word document.

    Args:
        file_path (str): The path of the Word document.

    Returns:
        dict: A dictionary with success status and message.
    """
    return word_add_page_break(file_path=file_path)


@mcp.tool()
def _word_search_and_replace(file_path: str, search_text: str, replace_text: str) -> dict:
    """
    Search and replace text in the Word document.

    Args:
        file_path (str): The path of the Word document.
        search_text (str): The text to search for.
        replace_text (str): The text to replace with.

    Returns:
        dict: A dictionary with success status and message.
    """
    return word_search_and_replace(file_path=file_path, search_text=search_text, replace_text=replace_text)


@mcp.tool()
def _word_format_text(file_path: str, text: str, font_name: Optional[str]='Times New Roman', font_size: Optional[int]=12, bold: Optional[bool]=False, italic: Optional[bool]=False, underline: Optional[bool]=False, color: Optional[str]='000000') -> dict:
    """
    Format text in the Word document.

    Args:
        file_path (str): The path of the Word document.
        text (str): The text to format.
        font_name (str): The font name. Default is "Times New Roman".
        font_size (int): The font size. Default is 12.
        bold (bool): Whether to make the text bold. Default is False.
        italic (bool): Whether to make the text italic. Default is False.
        underline (bool): Whether to underline the text. Default is False.
        color (str): The color of the text in hex format. Default is "000000".

    Returns:
        dict: A dictionary with success status and message.
    """
    return word_format_text(file_path=file_path, text=text, font_name=font_name, font_size=font_size, bold=bold, italic=italic, underline=underline, color=color)


@mcp.tool()
def _word_get_document_info(file_path: str) -> dict:
    """
    Get information about the Word document.

    Args:
        file_path (str): The path of the Word document.

    Returns:
        dict: A dictionary with success status and document information.
    """
    return word_get_document_info(file_path=file_path)


@mcp.tool()
def _word_get_document_text(file_path: str) -> dict:
    """
    Get the text content of the Word document.

    Args:
        file_path (str): The path of the Word document.

    Returns:
        dict: A dictionary with success status and document text.
    """
    return word_get_document_text(file_path=file_path)


@mcp.tool()
def _word_get_document_outline(file_path: str) -> dict:
    """
    Get the outline of the Word document.

    Args:
        file_path (str): The path of the Word document.

    Returns:
        dict: A dictionary with success status and document outline.
    """
    return word_get_document_outline(file_path=file_path)


@mcp.tool()
def _word_get_document_paragraph(file_path: str, index: int) -> dict:
    """
    Get a specific paragraph from the Word document.

    Args:
        file_path (str): The path of the Word document.
        index (int): The index of the paragraph.

    Returns:
        dict: A dictionary with success status and paragraph text.
    """
    return word_get_document_paragraph(file_path=file_path, index=index)


@mcp.tool()
def _word_copy_document(source_file_path: str, destination_file_path: str) -> dict:
    """
    Copy a Word document.

    Args:
        source_file_path (str): The path of the source Word document.
        destination_file_path (str): The path of the destination Word document.

    Returns:
        dict: A dictionary with success status and message.
    """
    return word_copy_document(source_file_path=source_file_path, destination_file_path=destination_file_path)


@mcp.tool()
def _word_list_documents(directory: str) -> dict:
    """
    List all Word documents in a directory.

    Args:
        directory (str): The path of the directory.

    Returns:
        dict: A dictionary with success status and list of document paths.
    """
    return word_list_documents(directory=directory)


if __name__ == '__main__':
    mcp.run('stdio')
